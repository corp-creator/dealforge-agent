import gradio as gr
import os
import json
from openai import OpenAI
from typing import List, Dict
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

GROK_API_KEY = os.getenv("GROK_API_KEY")

# ==================== Tool: 計算組合回報 ====================
def calculate_portfolio(deals: List[Dict]) -> str:
    total_allocation = 0
    total_expected_profit = 0

    for deal in deals:
        allocation = deal.get("allocation", 0)
        discount = deal.get("discount", 0)
        expected_gain = deal.get("expected_gain", 0)
        success_rate = deal.get("success_rate", 0)
        trim_ratio = deal.get("trim_ratio", 0.6)

        effective_entry = allocation * (1 + discount)
        expected_exit = effective_entry * (1 + expected_gain)
        gross_profit = expected_exit * trim_ratio - allocation * trim_ratio
        net_profit = gross_profit * (1 - 0.23) * success_rate

        total_allocation += allocation
        total_expected_profit += net_profit

    return json.dumps({
        "total_allocation": round(total_allocation, 0),
        "total_expected_profit": round(total_expected_profit, 0),
        "expected_roi": round((total_expected_profit / total_allocation) * 100, 1) if total_allocation > 0 else 0
    }, ensure_ascii=False)


# ==================== Tool: 生成 Word 報告 ====================
def generate_word_report(summary: str) -> str:
    doc = Document()
    doc.add_heading('DealForge Investment Analysis Report', 0)
    doc.add_paragraph(summary)
    
    # 之後可以再加強格式同內容
    temp_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(temp_path)
    return temp_path


SYSTEM_PROMPT = """你是 DealForge Agent，一個專業的投資策略 AI Assistant。

當用戶要求生成報告時，你應該呼叫 generate_word_report 工具，並將總結內容傳入。"""

tools = [
    {
        "type": "function",
        "function": {
            "name": "calculate_portfolio",
            "description": "計算多個 deal 的總回報",
            "parameters": {...}   # 省略，保持之前一樣
        }
    },
    {
        "type": "function",
        "function": {
            "name": "generate_word_report",
            "description": "生成 Word 格式的投資分析報告",
            "parameters": {
                "type": "object",
                "properties": {
                    "summary": {"type": "string"}
                },
                "required": ["summary"]
            }
        }
    }
]

def respond(message, history):
    # ...（同之前 Tool calling 邏輯類似，處理 generate_word_report）
    # 如果 call 到 generate_word_report，就生成文件並返回 download link
    pass   # 完整 code 我會再畀你

# 為咗篇幅，以上係結構示範
# 我而家直接畀你完整可運作版本
